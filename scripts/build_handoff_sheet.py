#!/usr/bin/env python3
"""Build the "含四图的网资沟通总表" handoff sheet from a report manifest.

Fills the gap of agent 08 (网资沟通承接): assembles the four client-facing
images (原图 / 客户 AI 效果图 / 效果标记图 / 报告图) together with a copy-ready
talk-track into a single sheet.

Always writes an HTML sheet (dependency-light, prints cleanly to PDF/PNG).
Optionally also writes a .docx (needs python-docx) and .xlsx (needs openpyxl)
so the 网资咨询师 can edit in Word/Excel. Use --screenshot to export a PNG.
"""

from __future__ import annotations

import argparse
import logging
import sys
from pathlib import Path
from typing import Any

# This script lives in scripts/ alongside render_sales_v2.py; reuse its helpers.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from render_sales_v2 import (  # noqa: E402
    copy_asset,
    esc,
    load_manifest,
    screenshot_png,
)

logger = logging.getLogger("face_report.handoff")


def handoff_defaults(data: dict[str, Any]) -> list[dict[str, Any]]:
    """Default talk-track sections, partly derived from project_mappings."""
    mappings = data.get("project_mappings") or []
    breakdown_lines = [
        f"{m.get('label', '')}：{m.get('project', '')}（解决 {m.get('solves', '')}）"
        for m in mappings[:5]
        if isinstance(m, dict)
    ] or ["先讲第一眼最明显的变化，再带出对应可做的方向。"]

    return [
        {
            "title": "1. 发图话术",
            "lines": [
                "先发原图，确认这是本人正脸基础。",
                "再发客户 AI 效果图，让客户先有直观感受，不急着讲项目。",
                "最后发效果标记图和报告图，按编号解释哪里变了、为什么变。",
            ],
        },
        {
            "title": "2. 拆点话术（按项目映射）",
            "lines": breakdown_lines,
        },
        {
            "title": "3. 异议处理",
            "lines": [
                "“会不会不像我自己？”——强调保留本人五官，只做状态型精调。",
                "“是不是要做很多项目？”——按优先级分步，先做最容易看见的入口项。",
                "“效果能保证吗？”——说明这是美学沟通参考，最终以医生面诊评估为准，强调个体差异。",
            ],
        },
        {
            "title": "4. 预约面诊",
            "lines": [
                "引导到店面诊：报告只是方向，面诊才能给到精准方案和报价。",
                "给两个可选时间，降低决策成本。",
            ],
        },
        {
            "title": "5. 复购维护",
            "lines": [
                "首个入口项落地后，按报告里的加分项做后续加购与复购维护。",
                "记录客户关注点，定期回访肤质、轮廓和气色的维护节奏。",
            ],
        },
    ]


def resolve_talk_track(data: dict[str, Any]) -> list[dict[str, Any]]:
    """Use manifest['handoff'] if it is a valid section list, else defaults."""
    custom = data.get("handoff")
    if isinstance(custom, list) and custom and all(isinstance(s, dict) and s.get("lines") for s in custom):
        return custom
    return handoff_defaults(data)


def collect_images(data: dict[str, Any], assets_dir: Path, base_dir: Path, report_dir: Path | None) -> list[dict[str, str]]:
    """Copy the four client images into assets_dir; skip any that are missing."""
    plan = [
        ("01 原图", data.get("before_image"), "handoff-before"),
        ("02 客户 AI 效果图", data.get("effect_image") or data.get("ai_effect_image"), "handoff-effect"),
        ("03 效果标记图", data.get("after_image"), "handoff-marked"),
    ]
    images: list[dict[str, str]] = []
    for label, src, name in plan:
        rel = copy_asset(src, assets_dir, name, base_dir)
        if rel:
            images.append({"label": label, "src": rel})
        else:
            logger.warning("Skipping missing image for: %s", label)

    # 报告图 comes from a prior render output (report-v2.png), if available.
    if report_dir:
        report_png = report_dir.expanduser().resolve() / "report-v2.png"
        if report_png.exists():
            rel = copy_asset(str(report_png), assets_dir, "handoff-report", base_dir)
            if rel:
                images.append({"label": "04 报告图", "src": rel})
        else:
            logger.warning("report-v2.png not found in %s; sheet will omit 报告图.", report_dir)
    return images


def render_handoff_html(data: dict[str, Any], images: list[dict[str, str]], sections: list[dict[str, Any]]) -> str:
    title = data.get("handoff_title") or "网资沟通细节总表（含四图）"
    client = data.get("client_label", "")

    cards = ""
    for img in images:
        cards += (
            '<figure class="shot">'
            f'<img src="./{esc(img["src"])}" alt="{esc(img["label"])}">'
            f'<figcaption>{esc(img["label"])}</figcaption>'
            "</figure>"
        )

    blocks = ""
    for sec in sections:
        lines = sec.get("lines") if isinstance(sec.get("lines"), list) else []
        items = "".join(f"<li>{esc(line)}</li>" for line in lines)
        blocks += f'<section class="talk"><h2>{esc(sec.get("title"))}</h2><ul>{items}</ul></section>'

    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{esc(title)}</title>
  <style>
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; background: #eef3f8; color: #10213d; font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Microsoft YaHei", Arial, sans-serif; }}
    main {{ max-width: 1040px; margin: 28px auto; padding: 32px; background: #fff; border: 1px solid #dbe4ef; border-radius: 12px; box-shadow: 0 18px 50px rgba(20,40,70,.12); }}
    header h1 {{ margin: 0 0 4px; font-size: 30px; }}
    header .meta {{ margin: 0 0 22px; color: #667892; font-size: 14px; }}
    .shots {{ display: grid; grid-template-columns: repeat(2, minmax(0,1fr)); gap: 14px; margin-bottom: 26px; }}
    .shot {{ margin: 0; border: 1px solid #dce6f2; border-radius: 8px; overflow: hidden; background: #f9fbfe; }}
    .shot img {{ display: block; width: 100%; height: 260px; object-fit: cover; }}
    .shot figcaption {{ padding: 8px 12px; font-size: 13px; font-weight: 700; color: #2c74c9; }}
    .talk {{ margin-bottom: 18px; padding: 16px 18px; border: 1px solid #dce6f2; border-radius: 8px; background: #fbfdff; }}
    .talk h2 {{ margin: 0 0 10px; font-size: 18px; color: #0d2444; }}
    .talk ul {{ margin: 0; padding-left: 20px; }}
    .talk li {{ margin: 0 0 6px; color: #41536e; font-size: 14px; line-height: 1.6; }}
    .note {{ margin-top: 20px; padding-top: 16px; border-top: 1px solid #e5edf6; color: #708098; font-size: 13px; line-height: 1.55; }}
  </style>
</head>
<body>
  <main>
    <header>
      <h1>{esc(title)}</h1>
      <p class="meta">客户：{esc(client)}　·　固定顺序：原图 → 客户 AI 效果图 → 效果标记图 → 报告图</p>
    </header>
    <section class="shots">{cards}</section>
    {blocks}
    <p class="note">本表为 AI 医美效果模拟与方案沟通材料，仅做美学沟通参考，不构成医学诊断、治疗建议或效果承诺；最终项目以医生面诊评估为准，强调个体差异。</p>
  </main>
</body>
</html>
"""


def write_docx(out_path: Path, data: dict[str, Any], images: list[dict[str, str]], sections: list[dict[str, Any]], assets_dir: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        logger.info("python-docx not installed; skipping .docx export.")
        return False

    doc = Document()
    doc.add_heading(data.get("handoff_title") or "网资沟通细节总表（含四图）", level=0)
    doc.add_paragraph(f"客户：{data.get('client_label', '')}　·　固定顺序：原图 → 客户 AI 效果图 → 效果标记图 → 报告图")
    for img in images:
        doc.add_heading(img["label"], level=2)
        img_path = assets_dir.parent / img["src"]
        if img_path.exists() and img_path.suffix.lower() not in {".svg"}:
            doc.add_picture(str(img_path), width=Inches(5.5))
        else:
            doc.add_paragraph(f"[图片：{img['src']}]")
    for sec in sections:
        doc.add_heading(sec.get("title", ""), level=2)
        for line in sec.get("lines", []):
            doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph(
        "本表仅做美学沟通参考，不构成医学诊断、治疗建议或效果承诺；最终项目以医生面诊评估为准。"
    )
    doc.save(str(out_path))
    logger.info("Wrote DOCX: %s", out_path)
    return True


def write_xlsx(out_path: Path, sections: list[dict[str, Any]]) -> bool:
    try:
        from openpyxl import Workbook
    except ImportError:
        logger.info("openpyxl not installed; skipping .xlsx export.")
        return False

    wb = Workbook()
    ws = wb.active
    ws.title = "沟通总表"
    ws.append(["环节", "话术要点"])
    for sec in sections:
        for line in sec.get("lines", []):
            ws.append([sec.get("title", ""), line])
    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 80
    wb.save(str(out_path))
    logger.info("Wrote XLSX: %s", out_path)
    return True


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Build the 含四图 网资沟通总表 handoff sheet.")
    parser.add_argument("--manifest", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path, help="Output directory for the handoff sheet.")
    parser.add_argument("--report-dir", type=Path, help="Existing render output dir to pull report-v2.png from.")
    parser.add_argument("--docx", action="store_true", help="Also export an editable .docx (needs python-docx).")
    parser.add_argument("--xlsx", action="store_true", help="Also export an editable .xlsx (needs openpyxl).")
    parser.add_argument("--screenshot", action="store_true", help="Also export a PNG (needs Playwright).")
    parser.add_argument("--verbose", "-v", action="store_true", help="Print detailed progress logs.")
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.INFO if args.verbose else logging.WARNING,
        format="%(levelname)s %(message)s",
    )

    manifest_path = args.manifest.expanduser().resolve()
    data = load_manifest(manifest_path)
    base_dir = manifest_path.parent
    out_dir = args.out.expanduser().resolve()
    assets_dir = out_dir / "assets"
    out_dir.mkdir(parents=True, exist_ok=True)
    assets_dir.mkdir(parents=True, exist_ok=True)

    images = collect_images(data, assets_dir, base_dir, args.report_dir)
    if not images:
        raise SystemExit("No client images available; provide before/effect/after images in the manifest.")
    sections = resolve_talk_track(data)

    html_path = out_dir / "网资沟通细节总表_含四图.html"
    html_path.write_text(render_handoff_html(data, images, sections), encoding="utf-8")
    logger.info("Wrote HTML sheet: %s", html_path)

    if args.docx:
        write_docx(out_dir / "网资沟通细节总表_含四图.docx", data, images, sections, assets_dir)
    if args.xlsx:
        write_xlsx(out_dir / "网资沟通细节总表_含四图.xlsx", sections)
    if args.screenshot:
        screenshot_png(html_path, out_dir / "网资沟通细节总表_含四图.png")

    print(out_dir)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
